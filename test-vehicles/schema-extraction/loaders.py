"""
Pre-treatment. Turns whatever a person handed you into the plain text that goes
into the API call. Nothing here talks to Claude.

The contract: whatever this returns IS the source of truth. `measure()` checks
the model's verbatim quotes against this text, so if extraction mangles the
document, coverage measures the mangling and not the schema. Hence the warnings —
they exist to stop you drawing conclusions from a bad read.
"""

from __future__ import annotations

from pathlib import Path

ALLOWED_FILE_TYPES = {".txt", ".md", ".pdf", ".docx"}

# First bytes of the file, which is what it actually IS regardless of what the
# extension claims. .docx is a zip, hence PK.
MAGIC = [
    (b"%PDF-", ".pdf"),
    (b"PK\x03\x04", ".docx"),
    (b"\xd0\xcf\x11\xe0", ".doc"),   # legacy OLE — Word 97-2003, not supported
]


class LoadError(Exception):
    """Raised when the file can't become text. Caught by the caller, not fatal."""


# ---------------------------------------------------------------------------
# detector
# ---------------------------------------------------------------------------

def detect(path: Path) -> str:
    """Returns the real extension, sniffed from content where possible.

    A resume emailed as `resume.pdf` that is secretly a .docx is a real thing
    that happens, and pdfplumber's error for it is inscrutable. Sniff first.
    """
    with open(path, "rb") as fh:
        head = fh.read(8)

    for sig, kind in MAGIC:
        if head.startswith(sig):
            return kind

    # No magic number: assume it's the text-ish thing the extension claims.
    suffix = path.suffix.lower()
    return suffix if suffix in ALLOWED_FILE_TYPES else ".txt"


# ---------------------------------------------------------------------------
# per-format readers
# ---------------------------------------------------------------------------

def _load_txt(path: Path) -> tuple[str, list[str]]:
    try:
        return path.read_text(encoding="utf-8"), []
    except UnicodeDecodeError:
        # Windows-authored .txt is often cp1252. Smart quotes, mostly.
        return path.read_text(encoding="cp1252", errors="replace"), [
            "not valid UTF-8; read as cp1252, a few characters may be wrong"
        ]


def _load_pdf(path: Path) -> tuple[str, list[str]]:
    import pdfplumber

    warnings: list[str] = []
    pages: list[str] = []

    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        for page in pdf.pages:
            pages.append(page.extract_text() or "")

    text = "\n\n".join(pages).strip()

    if not text:
        raise LoadError(
            "no text layer at all — this is a scan or an image export. "
            "It needs OCR, which this rig doesn't do. Ask for the .docx."
        )
    if len(text) / max(n_pages, 1) < 250:
        warnings.append(
            f"only {len(text) // max(n_pages, 1)} chars/page — probably a partial "
            f"scan or a heavily graphical layout. Eyeball it before trusting the run."
        )
    if n_pages > 4:
        warnings.append(f"{n_pages} pages — long for a resume, check it isn't a portfolio")

    return text, warnings


def _load_docx(path: Path) -> tuple[str, list[str]]:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(path)
    warnings: list[str] = []
    chunks: list[str] = []
    n_tables = 0

    # Walking body children rather than doc.paragraphs, because doc.paragraphs
    # silently omits everything inside tables — and a startling number of resumes
    # are one big invisible two-column table. That failure mode looks like an
    # empty document, not like an error.
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            line = Paragraph(child, doc).text.strip()
            if line:
                chunks.append(line)
        elif tag == "tbl":
            n_tables += 1
            for row in Table(child, doc).rows:
                cells = [c.text.strip() for c in row.cells]
                # de-dupe: merged cells report the same text more than once
                seen, kept = set(), []
                for c in cells:
                    if c and c not in seen:
                        seen.add(c)
                        kept.append(c)
                if kept:
                    chunks.append("  ".join(kept))

    text = "\n".join(chunks).strip()

    if not text:
        raise LoadError("opened fine but contained no text — check it in Word")
    if n_tables:
        warnings.append(
            f"{n_tables} table(s) flattened into lines — reading order may differ "
            f"from what the page looks like"
        )

    return text, warnings


READERS = {
    ".txt": _load_txt,
    ".md": _load_txt,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


# ---------------------------------------------------------------------------
# the one function lab.py calls
# ---------------------------------------------------------------------------

def load(path: Path) -> tuple[str, list[str]]:
    """Returns (text, warnings). Raises LoadError if there's no usable text."""
    real = detect(path)
    claimed = path.suffix.lower()
    warnings: list[str] = []

    if real == ".doc":
        raise LoadError("legacy .doc — open it in Word and Save As .docx")
    if real not in ALLOWED_FILE_TYPES:
        raise LoadError(f"unsupported type {real}")
    if real != claimed and claimed in ALLOWED_FILE_TYPES:
        warnings.append(f"named {claimed} but is actually {real}; read as {real}")

    text, more = READERS[real](path)
    return text, warnings + more
