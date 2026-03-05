"""
services/docx_builder.py

Dumb renderer — takes Claude's structured JSON and produces a .docx file.

DESIGN (ADR-011):
    All formatting decisions live in the prompt template, not here. Claude's
    JSON specifies font sizes, bold/italic ranges, hyperlinks, and element
    ordering. This module just walks the array and translates each element
    into python-docx calls. Change the prompt, change the output.

    The test: if a user changes their prompt to say "use 12pt for everything
    and never bold anything," this code should honor it without modification.

HARDCODED (document-level, not content-level):
    - Page size: US Letter (8.5 x 11)
    - Margins: 0.7" left/right, 0.5" top/bottom (tight for resumes)
    - Default font family: Calibri
    - These become configurable via a document-settings template in Phase N.

ELEMENT TYPES (the JSON contract between Claude and this renderer):
    contact_name    — centered name, specified font_size
    contact_info    — centered details, optional hyperlinks[]
    section_header  — uppercase text + paragraph bottom border
    paragraph       — body text with optional bold[]/italic[] substring arrays
    job_title       — position title line
    job_meta        — company / location / dates line
    bullet          — bulleted list item with bold[]/italic[]
    blank_line      — empty paragraph for intentional vertical spacing
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Document-level defaults (hardcoded — see ADR-011 note above) ─────────────

_DEFAULT_FONT = "Calibri"
_PAGE_WIDTH_INCHES = 8.5
_PAGE_HEIGHT_INCHES = 11.0
_MARGIN_TOP = 0.5
_MARGIN_BOTTOM = 0.5
_MARGIN_LEFT = 0.7
_MARGIN_RIGHT = 0.7


# ── Public API ────────────────────────────────────────────────────────────────

def build_resume_docx(elements: list[dict]) -> bytes:
    """
    Render a list of structured elements into a .docx byte string.

    Each element in the list is a dict with at minimum a "type" key.
    The renderer makes no formatting decisions — it executes whatever
    Claude put in the JSON. See module docstring for the element schema.

    Returns bytes suitable for storing in TailoringJob.output_resume_docx.
    """
    doc = Document()
    _configure_document(doc)

    for element in elements:
        _render_element(doc, element)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Document setup ────────────────────────────────────────────────────────────

def _configure_document(doc: Document) -> None:
    """Set page size, margins, and default font. Hardcoded for Phase 0."""
    style = doc.styles["Normal"]
    style.font.name = _DEFAULT_FONT
    style.font.size = Pt(11)

    for section in doc.sections:
        section.page_width = Inches(_PAGE_WIDTH_INCHES)
        section.page_height = Inches(_PAGE_HEIGHT_INCHES)
        section.top_margin = Inches(_MARGIN_TOP)
        section.bottom_margin = Inches(_MARGIN_BOTTOM)
        section.left_margin = Inches(_MARGIN_LEFT)
        section.right_margin = Inches(_MARGIN_RIGHT)


# ── Element dispatcher ────────────────────────────────────────────────────────

def _render_element(doc: Document, element: dict) -> None:
    """Dispatch a single element to its renderer by type."""
    el_type = element.get("type", "")
    renderers = {
        "contact_name": _render_contact_name,
        "contact_info": _render_contact_info,
        "section_header": _render_section_header,
        "paragraph": _render_paragraph,
        "job_title": _render_job_title,
        "job_meta": _render_job_meta,
        "bullet": _render_bullet,
        "blank_line": _render_blank_line,
    }
    renderer = renderers.get(el_type)
    if renderer:
        renderer(doc, element)
    # Unknown types are silently skipped — the prompt might evolve faster
    # than the renderer, and partial output is better than a crash.


# ── Element renderers ─────────────────────────────────────────────────────────
# Each function adds one or more paragraphs to the document. Font sizes,
# bold/italic ranges, and text content all come from the element dict.
# The renderer never invents formatting — it only executes what's specified.

def _render_contact_name(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(el.get("text", ""))
    run.font.size = Pt(el.get("font_size", 14))
    run.font.name = _DEFAULT_FONT
    run.bold = True


def _render_contact_info(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    text = el.get("text", "")
    hyperlinks = el.get("hyperlinks", [])
    font_size = el.get("font_size", 10)

    if not hyperlinks:
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = _DEFAULT_FONT
        return

    # Build the line with hyperlinks replacing their matching text segments
    remaining = text
    for link in hyperlinks:
        link_text = link.get("text", "")
        link_url = link.get("url", "")
        if link_text not in remaining:
            continue

        before, _, after = remaining.partition(link_text)
        if before:
            run = p.add_run(before)
            run.font.size = Pt(font_size)
            run.font.name = _DEFAULT_FONT

        _add_hyperlink(p, link_url, link_text, font_size)
        remaining = after

    if remaining:
        run = p.add_run(remaining)
        run.font.size = Pt(font_size)
        run.font.name = _DEFAULT_FONT


def _render_section_header(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(el.get("text", "").upper())
    run.font.size = Pt(el.get("font_size", 11))
    run.font.name = _DEFAULT_FONT
    # No bold — "plain uppercase, the underline bar does the structural work"
    run.bold = False

    _add_bottom_border(p)


def _render_paragraph(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _add_formatted_runs(p, el)


def _render_job_title(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(el.get("text", ""))
    run.font.size = Pt(el.get("font_size", 11))
    run.font.name = _DEFAULT_FONT
    run.bold = True


def _render_job_meta(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(el.get("text", ""))
    run.font.size = Pt(el.get("font_size", 10))
    run.font.name = _DEFAULT_FONT
    run.italic = True


def _render_bullet(doc: Document, el: dict) -> None:
    # python-docx's add_paragraph with style='List Bullet' uses the built-in
    # bullet style. This is cleaner than manual numbering config and renders
    # correctly in Word. (The SKILL.md warning about unicode bullets applies
    # to docx-js, not python-docx.)
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Clear the auto-generated run (style adds one), build our own with formatting
    p.clear()
    _add_formatted_runs(p, el)


def _render_blank_line(doc: Document, _el: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Truly empty — the blank line IS the content (intentional spacing)


# ── Inline formatting ─────────────────────────────────────────────────────────

def _add_formatted_runs(paragraph, el: dict) -> None:
    """
    Add text to a paragraph with inline bold/italic applied to specific
    substrings. Claude specifies which substrings to format; this function
    splits the text at those boundaries and creates separate runs.

    Does not handle overlapping bold+italic on the same substring (rare edge
    case — would need interval merging). If a substring appears in both
    bold[] and italic[], bold wins. Good enough for Phase 0.
    """
    text = el.get("text", "")
    font_size = el.get("font_size", 11)
    bold_ranges = el.get("bold", [])
    italic_ranges = el.get("italic", [])

    if not bold_ranges and not italic_ranges:
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = _DEFAULT_FONT
        return

    # Build intervals: (start_index, end_index, format_type)
    intervals = []
    for substr in bold_ranges:
        idx = text.find(substr)
        if idx >= 0:
            intervals.append((idx, idx + len(substr), "bold"))
    for substr in italic_ranges:
        idx = text.find(substr)
        if idx >= 0:
            intervals.append((idx, idx + len(substr), "italic"))

    intervals.sort(key=lambda x: x[0])

    # Walk text, emitting runs with appropriate formatting
    pos = 0
    for start, end, fmt in intervals:
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.font.size = Pt(font_size)
            run.font.name = _DEFAULT_FONT
        run = paragraph.add_run(text[start:end])
        run.font.size = Pt(font_size)
        run.font.name = _DEFAULT_FONT
        if fmt == "bold":
            run.bold = True
        elif fmt == "italic":
            run.italic = True
        pos = end

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(font_size)
        run.font.name = _DEFAULT_FONT


# ── XML helpers (python-docx doesn't expose these natively) ───────────────────

def _add_bottom_border(paragraph) -> None:
    """
    Add a bottom border (paragraph underline) below a paragraph.
    Used for section headers — "the underline bar does the structural work."
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")       # border weight in half-points
    bottom.set(qn("w:space"), "1")    # space between text and border
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url: str, text: str, font_size: float) -> None:
    """
    Add a clickable hyperlink inline in a paragraph.
    python-docx doesn't have a high-level hyperlink API, so we build
    the XML directly.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Hyperlink styling: blue + underline
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Font size (half-points in OOXML)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(szCs)

    # Font family
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _DEFAULT_FONT)
    rFonts.set(qn("w:hAnsi"), _DEFAULT_FONT)
    rPr.append(rFonts)

    run_elem.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)
